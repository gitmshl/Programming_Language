from docx import Document
import mpp_exceptions


key_words = ["save", "+", "+A4", "+h", "+p"]


def handle_text(line, line_number, text, env):
    text = text[1:-1]
    result_text = []
    tokens = text.split()
    for token in tokens:
        if token == "$_":
            result_text.append(str(env["$_"]))
            continue
        token_lst = token.split('->')
        if token_lst[0] in env:
            var_name = token_lst[0]
            tmp = env[var_name]
            try:
                for tok in token_lst[1:]:
                    tmp = eval(f"tmp.{tok}")
            except:
                raise mpp_exceptions.VarAttributeError(line, line_number, token)
            result_text.append(tmp)
        else:
            result_text.append(token)
    try:
        ret = " ".join(result_text)
        return ret
    except:
        raise mpp_exceptions.VarTypeError(line, line_number)


def com_save(line_lst, line, line_number, env):
    if len(line_lst) not in (1, 2):
        raise mpp_exceptions.SyntaxError(line, line_number, "Неверное число параметров.")
    if len(line_lst) == 2:
        doc_name = line_lst[1]
    else:
        doc_name = "docx_document__default__"
    if doc_name not in env:
            raise mpp_exceptions.UnknownVar(line, line_number, doc_name)
    env[doc_name][0].save(env[doc_name][1])


def com_add_document(line_lst, line, line_number, env):
    def get_names():
        if len(line_lst) == 3:
            var_name, doc_name = line_lst[1], line_lst[2]
            if var_name in env:
                raise mpp_exceptions.DoubleVarUse(line, line_number)
        elif len(line_lst) == 2:
            var_name, doc_name = "docx_document__default__", line_lst[1]
        else:
            raise mpp_exceptions.SyntaxError(line, line_number, "Неверное число параметров.")
        if len(doc_name) < 6 or doc_name[-5:] != ".docx":
            raise mpp_exceptions.SyntaxError(line, line_number, "Неверное название файла docx")
        return var_name, doc_name
    
    var_name, doc_name = get_names()
    document = Document()
    env[var_name] = (document, doc_name)
    document.save(doc_name)


def com_add_A4(line_lst, line, line_number, env):
    if len(line_lst) > 2:
        raise mpp_exceptions.SyntaxError(line, line_number, "Неверное число параметров.")
    if len(line_lst) == 2:
        doc_name = line_lst[1]
        if doc_name not in env:
            raise mpp_exceptions.UnknownVar(line, line_number, doc_name)
    else:
        doc_name = "docx_document__default__"
    if env[doc_name][0] is None:
        raise mpp_exceptions.UseUnitiailizeVar(line, line_number, doc_name)

    env[doc_name][0].add_page_break()


def com_add_head(line_lst, line, line_number, env):
    if len(line_lst) < 3:
        raise mpp_exceptions.SyntaxError(line, line_number, "Неверное число параметров.")
    
    if len(line_lst) >= 4 and line_lst[2][0] == '"':
        doc_name, text, level = line_lst[1], " ".join(line_lst[2:-1]), line_lst[-1]
    else:
        doc_name, text, level = "docx_document__default__", " ".join(line_lst[1:-1]), line_lst[-1]

    if doc_name not in env:
            raise mpp_exceptions.UnknownVar(line, line_number, doc_name)
    try:
        level = int(level)
        if not (level >= 0 and level <= 9):
            raise ""
    except:
        raise mpp_exceptions.SyntaxError(line, line_number, "Некорректный параметр уровня level.")

    text = handle_text(line, line_number, text, env)
    env[doc_name][0].add_heading(text, level=level)


def com_add_paragraph(line_lst, line, line_number, env):
    if len(line_lst) < 2:
        raise mpp_exceptions.SyntaxError(line, line_number, "Неверное число параметров.")

    if len(line_lst) >= 3 and line_lst[2][0] == '"':
        doc_name, text = line_lst[1], " ".join(line_lst[2:])
    else:
        doc_name, text = "docx_document__default__", " ".join(line_lst[1:])

    if doc_name not in env:
            raise mpp_exceptions.UnknownVar(line, line_number, doc_name)
    
    text = handle_text(line, line_number, text, env)
    env[doc_name][0].add_paragraph(text)


def interpret_simple_command(line, line_number, env):
    line_lst = line.split() if line[0] != '[' else line.split()[1:]
    command = line_lst[0]
    if command not in key_words:
        raise mpp_exceptions.UnknownCommand(line, line_number)
    if command == "save":
        com_save(line_lst, line, line_number, env)
    elif command == "+":
        com_add_document(line_lst, line, line_number, env)
    elif command == "+A4":
        com_add_A4(line_lst, line, line_number, env)
    elif command == "+h":
        com_add_head(line_lst, line, line_number, env)
    elif command == "+p":
        com_add_paragraph(line_lst, line, line_number, env)
    else:
        print(f"unknow command:{line}")    
